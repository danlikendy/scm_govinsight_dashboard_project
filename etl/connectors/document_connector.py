"""
Коннектор для парсинга документов (PDF, DOCX)
"""

import requests
import pandas as pd
import io
import hashlib
from typing import Dict, Any, Optional, List
import re
from datetime import datetime
from .base_connector import BaseConnector

try:
    import tabula
    import PyPDF2
    from docx import Document
    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False
    print("Warning: tabula-py, PyPDF2, or python-docx not available. Document parsing will be limited.")


class DocumentConnector(BaseConnector):
    """Коннектор для извлечения данных из документов (PDF, DOCX)"""
    
    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'SCM-Dashboard/1.0'
        })
    
    def extract(self, url: Optional[str] = None, **kwargs) -> pd.DataFrame:
        """Извлечение данных из документа"""
        target_url = url or self.config.get('endpoint', '')
        self.logger.info(f"Extracting data from document {target_url}")
        
        try:
            # Загрузка файла
            response = self.session.get(target_url, timeout=60)
            response.raise_for_status()
            
            content = response.content
            file_hash = hashlib.sha256(content).hexdigest()
            
            # Определение типа файла
            file_extension = target_url.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                df = self._extract_from_pdf(content)
            elif file_extension == 'docx':
                df = self._extract_from_docx(content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Добавление метаданных
            df = self.add_metadata(df)
            df['_file_hash'] = file_hash
            df['_file_type'] = file_extension
            
            # Валидация
            schema = self.config.get('schema', [])
            df = self.validate_data(df, schema)
            
            self.log_extraction_stats(df)
            return df
            
        except Exception as e:
            self.logger.error(f"Error extracting data from {target_url}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, content: bytes) -> pd.DataFrame:
        """Извлечение данных из PDF файла"""
        if not TABULA_AVAILABLE:
            raise ImportError("tabula-py is required for PDF parsing")
        
        try:
            # Попытка извлечь таблицы с помощью tabula
            tables = tabula.read_pdf(
                io.BytesIO(content),
                pages='all',
                multiple_tables=True
            )
            
            if tables:
                # Объединяем все таблицы
                df = pd.concat(tables, ignore_index=True)
                self.logger.info(f"Extracted {len(df)} rows from PDF tables")
                return df
            else:
                # Если таблиц нет, извлекаем текст и применяем правила
                return self._extract_from_pdf_text(content)
                
        except Exception as e:
            self.logger.warning(f"Tabula extraction failed: {str(e)}, trying text extraction")
            return self._extract_from_pdf_text(content)
    
    def _extract_from_pdf_text(self, content: bytes) -> pd.DataFrame:
        """Извлечение данных из текста PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            # Применение правил извлечения
            extraction_rules = self.config.get('parsing', {}).get('extraction_rules', [])
            data = self._apply_extraction_rules(text, extraction_rules)
            
            return pd.DataFrame(data)
            
        except Exception as e:
            self.logger.error(f"PDF text extraction failed: {str(e)}")
            return pd.DataFrame()
    
    def _extract_from_docx(self, content: bytes) -> pd.DataFrame:
        """Извлечение данных из DOCX файла"""
        try:
            doc = Document(io.BytesIO(content))
            
            # Извлечение таблиц
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    # Первая строка - заголовки
                    headers = table_data[0]
                    rows = table_data[1:]
                    
                    for row in rows:
                        if len(row) == len(headers):
                            row_dict = dict(zip(headers, row))
                            tables_data.append(row_dict)
            
            if tables_data:
                df = pd.DataFrame(tables_data)
                self.logger.info(f"Extracted {len(df)} rows from DOCX tables")
                return df
            else:
                # Если таблиц нет, извлекаем текст и применяем правила
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                extraction_rules = self.config.get('parsing', {}).get('extraction_rules', [])
                data = self._apply_extraction_rules(text, extraction_rules)
                return pd.DataFrame(data)
                
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {str(e)}")
            return pd.DataFrame()
    
    def _apply_extraction_rules(self, text: str, rules: List[Dict]) -> List[Dict]:
        """Применение правил извлечения к тексту"""
        data = []
        
        for rule in rules:
            pattern = rule.get('pattern', '')
            field = rule.get('field', '')
            format_str = rule.get('format', None)
            
            if pattern and field:
                matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
                
                for match in matches:
                    value = match if isinstance(match, str) else match[0] if match else None
                    
                    if value:
                        # Форматирование значения
                        if format_str and 'date' in field.lower():
                            try:
                                value = datetime.strptime(value, format_str).date()
                            except:
                                pass
                        
                        data.append({field: value})
        
        return data
    
    def close(self):
        """Закрытие сессии"""
        self.session.close()
